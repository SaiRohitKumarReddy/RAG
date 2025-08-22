import os
import sys
os.environ["STREAMLIT_SERVER_FILE_WATCHER_TYPE"] = "none"
os.environ["STREAMLIT_SERVER_ENABLE_FILE_WATCHER"] = "false"
os.environ["STREAMLIT_SERVER_ENABLE_STATIC_SERVING"] = "false"
os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'  # For compatibility

import asyncio
if sys.platform == "win32":
    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
elif hasattr(asyncio, 'DefaultEventLoopPolicy'):
    asyncio.set_event_loop_policy(asyncio.DefaultEventLoopPolicy())

import streamlit as st
st.set_page_config(
    page_title="Document Analyzer Suite",
    layout="wide",
    initial_sidebar_state="collapsed"
)

import tempfile
import io
import traceback
import numpy as np
import pandas as pd
from PIL import Image, ImageEnhance, ImageFilter
from docx import Document
import PyPDF2
import pdfplumber
import pytesseract
import re
from openai import OpenAI
import fitz  # PyMuPDF
import json
import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

# Langchain imports for Advanced Analyzer
from langchain.text_splitter import RecursiveCharacterTextSplitter
try:
    from langchain_community.vectorstores import FAISS
    import faiss
    FAISS_AVAILABLE = True
except ImportError:
    FAISS_AVAILABLE = False
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate

# Shared: OpenAI Client and API Key
api_key = st.secrets.get("OPENAI_API_KEY")
if not api_key:
    st.error("OPENAI_API_KEY not found in secrets! Add it to .streamlit/secrets.toml")
    st.stop()
client = OpenAI(api_key=api_key)

# Shared: Lazy load Torch
@st.cache_resource
def get_torch():
    import torch
    return torch

# Shared: Lazy load Tesseract
@st.cache_resource
def get_tesseract():
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        st.info("Tesseract OCR initialized successfully")
        return pytesseract, True
    except Exception as e:
        st.error(f"Tesseract initialization failed: {str(e)}. OCR functionality will be disabled.")
        return None, False

# Shared: Image Preprocessing for OCR
def preprocess_image_for_ocr(image):
    image = image.convert("L")  # Grayscale
    image = image.filter(ImageFilter.MedianFilter())
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(2)  # Increase contrast
    return image

def extract_text_with_ocr(image):
    pytesseract, available = get_tesseract()
    if not available:
        return ""
    try:
        processed_image = preprocess_image_for_ocr(image)
        config = "--psm 6"  # Assume a single uniform block of text for handwriting
        text = pytesseract.image_to_string(processed_image, lang='eng', config=config)
        return text.strip()
    except Exception as e:
        st.error(f"OCR Error: {str(e)}")
        return ""

# --- Excel/CSV Processing ---
def process_spreadsheet(file, file_type):
    try:
        if file_type == 'csv':
            df = pd.read_csv(file)
        elif file_type == 'xlsx':
            # Skip the first row if it's a title, and use the second row as headers
            df = pd.read_excel(file, engine='openpyxl', skiprows=1)
        else:
            return None, "Unsupported spreadsheet format"
        
        # Clean column names (remove unnamed if any)
        df.columns = [col if not str(col).startswith('Unnamed') else f'Column_{i}' for i, col in enumerate(df.columns)]
        
        # Basic analysis
        analysis = {
            'row_count': len(df),
            'column_count': len(df.columns),
            'columns': list(df.columns),
            'missing_values': df.isnull().sum().to_dict(),
            'numeric_columns': df.select_dtypes(include=[np.number]).columns.tolist(),
            'summary_stats': df.describe().to_dict() if not df.select_dtypes(include=[np.number]).empty else None
        }
        return df, analysis
    except Exception as e:
        return None, f"Error processing spreadsheet: {str(e)}"

# --- Export Functions ---
def export_to_word(content, filename):
    doc = docx.Document()
    doc.add_heading(filename, 0)
    doc.add_paragraph(content)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def export_to_pdf(content, filename):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter  # 612 x 792 points
    left_margin = 1 * inch
    right_margin = 1 * inch
    top_margin = 1 * inch
    bottom_margin = 1 * inch
    max_line_width = width - left_margin - right_margin
    font_size = 12
    leading = font_size * 1.2  # Line spacing

    c.setFont("Helvetica", font_size)
    
    # Start text object
    textobject = c.beginText(left_margin, height - top_margin)
    y = height - top_margin
    
    # Split content into lines and wrap long ones
    lines = content.split('\n')
    for line in lines:
        # Wrap long lines
        while len(line) > 0:
            # Find max characters that fit in width (approximate, or use exact width)
            wrapped_line = line
            while c.stringWidth(wrapped_line, "Helvetica", font_size) > max_line_width and len(wrapped_line) > 1:
                wrapped_line = wrapped_line[:-1]
            if wrapped_line:
                textobject.textLine(wrapped_line)
                y -= leading
            line = line[len(wrapped_line):].lstrip()  # Remaining part
            
            if y <= bottom_margin:
                # Draw current text and start new page
                c.drawText(textobject)
                c.showPage()
                c.setFont("Helvetica", font_size)
                textobject = c.beginText(left_margin, height - top_margin)
                y = height - top_margin

    # Draw any remaining text
    c.drawText(textobject)
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def export_to_json(content, filename):
    data = {'filename': filename, 'content': content}
    return json.dumps(data, indent=2).encode('utf-8')

def export_to_excel(df, filename):
    buffer = io.BytesIO()
    df.to_excel(buffer, index=False, engine='openpyxl')
    buffer.seek(0)
    return buffer.getvalue()

# --- Summarizer Mode Functions ---
def summarize_is_text_substantial(text, min_words=30):
    if not text or not text.strip():
        return False
    words = re.findall(r'\b\w+\b', text.lower())
    meaningful_words = [w for w in words if len(w) > 2]
    return len(meaningful_words) >= min_words

def summarize_find_index_pages(file, file_type):
    index_keywords = [
        'table of contents', 'contents', 'index', 'table of content',
        'chapter', 'section', 'outline', 'overview'
    ]
    potential_pages = []
    try:
        if file_type == 'pdf':
            file.seek(0)
            reader = PyPDF2.PdfReader(file)
            for page_num in range(min(10, len(reader.pages))):
                try:
                    page_text = reader.pages[page_num].extract_text().lower()
                    for keyword in index_keywords:
                        if keyword in page_text:
                            potential_pages.append(page_num)
                            break
                    if re.search(r'\d+\s*\.\s*\d+|\d+\s*-\s*\d+', page_text):
                        potential_pages.append(page_num)
                except Exception:
                    continue
        elif file_type == 'docx':
            file.seek(0)
            doc = Document(file)
            for para in doc.paragraphs[:50]:
                text = para.text.lower()
                for keyword in index_keywords:
                    if keyword in text:
                        potential_pages.append(0)
                        break
        elif file_type in ['jpg', 'png', 'tiff']:
            file.seek(0)
            image = Image.open(file)
            text = extract_text_with_ocr(image)
            for keyword in index_keywords:
                if keyword in text.lower():
                    potential_pages.append(0)
                    break
    except Exception as e:
        st.warning(f"Error while searching for index pages: {str(e)}")
    return list(set(potential_pages))

def summarize_extract_text_from_pdf_pages(file, page_numbers):
    combined_text = ""
    try:
        file.seek(0)
        reader = PyPDF2.PdfReader(file)
        for page_num in page_numbers:
            if page_num < len(reader.pages):
                page_text = reader.pages[page_num].extract_text() or ""
                if not page_text.strip():
                    try:
                        file.seek(0)
                        with pdfplumber.open(file) as pdf:
                            if page_num < len(pdf.pages):
                                page = pdf.pages[page_num]
                                page_text = page.extract_text() or ""
                                if not page_text.strip():
                                    pil_image = page.to_image(resolution=300).original
                                    page_text = extract_text_with_ocr(pil_image)
                    except Exception:
                        pass
                combined_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
    except Exception as e:
        st.warning(f"Error extracting text from multiple pages: {str(e)}")
    return combined_text.strip()

def summarize_extract_text_from_pdf(file):
    text = ""
    try:
        file.seek(0)
        reader = PyPDF2.PdfReader(file)
        if len(reader.pages) > 0:
            text = reader.pages[0].extract_text() or ""
    except Exception as e:
        st.warning(f"PyPDF2 extraction failed: {str(e)}")
    if not text.strip():
        try:
            file.seek(0)
            with pdfplumber.open(file) as pdf:
                if len(pdf.pages) > 0:
                    page = pdf.pages[0]
                    text = page.extract_text() or ""
                    if not text.strip():
                        pil_image = page.to_image(resolution=300).original
                        text = extract_text_with_ocr(pil_image)
        except Exception as e:
            st.warning(f"OCR extraction failed: {str(e)}")
    return text.strip()

def summarize_extract_text_from_docx(file):
    try:
        file.seek(0)
        doc = Document(file)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return full_text[:1500] if full_text else ""
    except Exception as e:
        st.error(f"DOCX extraction failed: {str(e)}")
        return ""

def summarize_extract_text_from_image(file):
    try:
        file.seek(0)
        image = Image.open(file)
        text = extract_text_with_ocr(image)
        return text
    except Exception as e:
        st.error(f"Image extraction failed: {str(e)}")
        return ""

def summarize_extract_text_smart(file, file_type):
    extraction_log = []
    extraction_log.append("Attempting to extract text from first page...")
    if file_type == 'pdf':
        first_page_text = summarize_extract_text_from_pdf(file)
    elif file_type == 'docx':
        first_page_text = summarize_extract_text_from_docx(file)
    elif file_type in ['jpg', 'png', 'tiff']:
        first_page_text = summarize_extract_text_from_image(file)
    else:
        return "", ["Unsupported file type"], "unsupported"
    
    if summarize_is_text_substantial(first_page_text):
        extraction_log.append("First page contains substantial content")
        return first_page_text, extraction_log, "first_page"

    extraction_log.append("First page has limited content, searching for index/table of contents...")
    index_pages = summarize_find_index_pages(file, file_type)
    if index_pages:
        extraction_log.append(f"Found potential index pages: {[p + 1 for p in index_pages]}")
        if file_type == 'pdf':
            index_text = summarize_extract_text_from_pdf_pages(file, index_pages[:3])
        elif file_type == 'docx':
            file.seek(0)
            doc = Document(file)
            full_text = "\n".join([para.text for para in doc.paragraphs[:100] if para.text.strip()])
            index_text = full_text[:3000]
        else:  # image
            index_text = summarize_extract_text_from_image(file)
        if summarize_is_text_substantial(index_text):
            extraction_log.append("Successfully extracted content from index pages")
            return index_text, extraction_log, "index_pages"
        else:
            extraction_log.append("Index pages also contain limited content")
    else:
        extraction_log.append("No index/table of contents found")

    extraction_log.append("Using first page content as fallback")
    return first_page_text, extraction_log, "first_page_fallback"

def summarize_text_with_openai(text, extraction_method):
    try:
        if extraction_method == "index_pages":
            system_prompt = """You are a helpful assistant that creates clear, concise summaries. 
            The text provided appears to be from a table of contents or index section. 
            Create a summary that captures the main topics and structure of the document based on this index information.
            Always format your response with each bullet point on a separate line using the format: - Bullet point text."""
            user_prompt = f"""Please analyze this table of contents/index and create a 4-6 bullet point summary:
            {text}"""
        else:
            system_prompt = """You are a helpful assistant that creates clear, concise summaries. 
            Always format your response with each bullet point on a separate line using the format: - Bullet point text."""
            user_prompt = f"""Please summarize the following text in 3-5 bullet points:
            {text}"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=500,
            temperature=0.3
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"API Error: {str(e)}"

# --- Spreadsheet AI Analysis Function ---
def analyze_spreadsheet_with_openai(df):
    try:
        # Convert DF to string for prompt
        df_str = df.to_csv(index=False)
        system_prompt = """You are a data analyst. Provide insights, key trends, and summaries from the spreadsheet data."""
        user_prompt = f"""Analyze this spreadsheet data and provide 3-5 key insights:
        {df_str}"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=500,
            temperature=0.3
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"API Error: {str(e)}"

# --- Analyzer Mode Functions ---
@st.cache_resource
def analyzer_get_pdf_reader():
    from PyPDF2 import PdfReader
    return PdfReader

@st.cache_resource
def analyzer_get_docx_reader():
    from docx import Document
    return Document

@st.cache_resource
def analyzer_get_text_splitter():
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    return RecursiveCharacterTextSplitter

@st.cache_resource
def analyzer_get_embeddings():
    from langchain_openai import OpenAIEmbeddings
    return OpenAIEmbeddings

@st.cache_resource
def analyzer_get_openai():
    from langchain_openai import ChatOpenAI
    from langchain.chains import RetrievalQA
    from langchain.prompts import PromptTemplate
    return ChatOpenAI, RetrievalQA, PromptTemplate

@st.cache_resource
def analyzer_get_fitz():
    try:
        import fitz
        return fitz, True
    except ImportError:
        return None, False

def analyzer_determine_document_type(file, file_type, use_ocr):
    if file_type == 'docx':
        return "Word Document"
    elif file_type in ['jpg', 'png', 'tiff']:
        return "Image Document (OCR required)"
    
    try:
        PdfReader = analyzer_get_pdf_reader()
        file.seek(0)
        pdf_reader = PdfReader(file)
        fitz, fitz_available = analyzer_get_fitz()
        
        file.seek(0)
        pdf_bytes = file.read()
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        has_images = False
        image_count = 0
        
        for page_num in range(min(len(pdf_document), 3)):
            page = pdf_document.load_page(page_num)
            images = page.get_images()
            if images:
                has_images = True
                image_count += len(images)
        
        pdf_document.close()
        
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text += page_text
        
        text_length = len(text.strip())
        
        if has_images and image_count > 5:
            return "Infographic/Visual PDF"
        elif has_images and text_length < 200:
            return "Scanned PDF (OCR recommended)"
        elif has_images and text_length > 200:
            return "Mixed Content PDF (Text + Images)"
        elif text_length > 100:
            return "Text-based PDF"
        elif text_length > 0:
            return "Sparse Text PDF"
        else:
            return "Image-only PDF (OCR required)"
            
    except Exception as e:
        st.warning(f"Could not determine document type: {str(e)}")
        return "Unknown"

def analyzer_extract_text_from_pdf(pdf_file):
    try:
        PdfReader = analyzer_get_pdf_reader()
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page_text
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return None

def analyzer_extract_text_from_docx(docx_file):
    try:
        Document = analyzer_get_docx_reader()
        docx_file.seek(0)
        doc = Document(docx_file)
        text = ""
        for para_num, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                text += f"\n--- Paragraph {para_num + 1} ---\n"
                text += paragraph.text
        return text
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        return None

def analyzer_extract_text_from_image(image_file):
    try:
        image_file.seek(0)
        image = Image.open(image_file)
        text = extract_text_with_ocr(image)
        return text
    except Exception as e:
        st.error(f"Image extraction failed: {str(e)}")
        return ""

def analyzer_extract_images_and_ocr(pdf_file):
    pytesseract, tesseract_available = get_tesseract()
    fitz, fitz_available = analyzer_get_fitz()
    
    if not fitz_available:
        st.warning("PDF image processing (PyMuPDF) not available. Install with: pip install pymupdf")
        return ""

    if not tesseract_available:
        st.warning("Tesseract OCR not available. Install pytesseract and Tesseract binary.")
        return ""

    try:
        pdf_file.seek(0)
        pdf_bytes = pdf_file.read()
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        ocr_text = ""

        for page_num in range(min(len(pdf_document), 10)):
            try:
                page = pdf_document.load_page(page_num)
                mat = fitz.Matrix(1.5, 1.5)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                image = Image.open(io.BytesIO(img_data))

                page_ocr_text = f"\n--- OCR Page {page_num + 1} ---\n"
                config = "--psm 6"  # Assume a single uniform block of text for handwriting
                page_ocr_text += pytesseract.image_to_string(image, lang='eng', config=config)
                ocr_text += page_ocr_text + "\n"
            except Exception as e:
                st.warning(f"OCR processing failed for page {page_num + 1}: {str(e)}")
                continue

        pdf_document.close()
        if ocr_text:
            st.info("OCR completed successfully")
        return ocr_text
    except Exception as e:
        st.error(f"OCR processing failed: {str(e)}. Continuing without OCR.")
        return ""

@st.cache_resource
def analyzer_create_embeddings():
    try:
        torch = get_torch()
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            
        OpenAIEmbeddings = analyzer_get_embeddings()
        embeddings = OpenAIEmbeddings(
            model="text-embedding-3-small",
            openai_api_key=api_key
        )
        
        test_embedding = embeddings.embed_query("test")
        if test_embedding and len(test_embedding) > 0:
            return embeddings
        else:
            raise ValueError("Empty embedding")
            
    except Exception as e:
        st.error(f"Embedding creation failed: {str(e)}")
        return None

def analyzer_create_vector_store(text_chunks):
    try:
        if not FAISS_AVAILABLE:
            st.error("FAISS not available. Install with: pip install faiss-cpu")
            return None

        valid_chunks = [
            ' '.join(chunk.strip().split()) 
            for chunk in text_chunks 
            if chunk and isinstance(chunk, str) and len(chunk.strip()) > 10
        ]
        
        if not valid_chunks:
            st.error("No valid text chunks")
            return None

        st.info(f"Processing {len(valid_chunks)} chunks...")
        embeddings = analyzer_create_embeddings()
        if not embeddings:
            return None
            
        batch_size = 20
        vector_store = None
        
        progress_bar = st.progress(0)
        for i in range(0, len(valid_chunks), batch_size):
            batch = valid_chunks[i:i + batch_size]
            try:
                if vector_store is None:
                    vector_store = FAISS.from_texts(batch, embedding=embeddings)
                else:
                    batch_vs = FAISS.from_texts(batch, embedding=embeddings)
                    vector_store.merge_from(batch_vs)
                
                progress = min((i + batch_size) / len(valid_chunks), 1.0)
                progress_bar.progress(progress)
            except Exception as e:
                st.warning(f"Batch {i//batch_size + 1} failed: {str(e)}")
                continue

        progress_bar.empty()
        
        if vector_store:
            st.success("Vector store created")
            return vector_store
        else:
            st.error("Vector store creation failed")
            return None
            
    except Exception as e:
        st.error(f"Vector store error: {str(e)}")
        return None

def analyzer_create_qa_chain(vector_store, api_key):
    try:
        ChatOpenAI, RetrievalQA, PromptTemplate = analyzer_get_open

ai()
        
        llm = ChatOpenAI(
            openai_api_key=api_key,
            model_name="gpt-4o-mini",
            temperature=0.3,
            max_tokens=1024
        )

        prompt_template = """
You are a document analyst. Answer based on the provided context.

Context: {context}

Question: {question}

Instructions:
- Provide detailed answers when context supports it
- For numerical data, present it clearly
- If information is incomplete, say what you found
- If no relevant info exists, state "Information not found in document"

Answer:
"""

        prompt = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )

        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vector_store.as_retriever(search_kwargs={"k": 4}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": prompt}
        )
        return qa_chain
    except Exception as e:
        st.error(f"QA chain error: {str(e)}")
        return None

# --- Main App Logic ---
def main():
    st.title("Document Analyzer Suite")
    st.markdown("Choose a mode to analyze your PDF, Word, image, or spreadsheet.")

    # Mode Selection
    mode = st.radio(
        "Select Mode:",
        ["Summarize the Document", "Advanced Analyzer"],
        horizontal=True
    )

    # Shared System Info
    with st.expander("System Info"):
        torch = get_torch()
        st.info(f"PyTorch: {torch.__version__}")
        st.info(f"CUDA: {torch.cuda.is_available()}")

    if mode == "Summarize the Document":
        st.header("Document Summarizer")
        st.markdown("*AI-powered summarization with smart content detection*")

        uploaded_file = st.file_uploader("Choose a file to summarize", type=["pdf", "docx", "jpg", "png", "tiff", "csv", "xlsx"], key="summarize_uploader")

        if uploaded_file is not None:
            file_extension = uploaded_file.name.split(".")[-1].lower()
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("File Name", uploaded_file.name)
            with col2:
                st.metric("File Size", f"{uploaded_file.size / 1024:.1f} KB")
            with col3:
                st.metric("File Type", file_extension.upper())

            if file_extension in ['csv', 'xlsx']:
                with st.spinner("Analyzing spreadsheet..."):
                    df, analysis = process_spreadsheet(uploaded_file, file_extension)

                if df is not None:
                    st.subheader("Spreadsheet Preview")
                    st.dataframe(df.head())

                    st.subheader("Basic Analysis Summary")
                    st.markdown(f"""
                    - **Rows**: {analysis['row_count']}
                    - **Columns**: {analysis['column_count']}
                    - **Column Names**: {', '.join(analysis['columns'])}
                    - **Missing Values**: {analysis['missing_values']}
                    """)
                    if analysis['summary_stats']:
                        st.markdown("**Summary Statistics**:")
                        st.dataframe(pd.DataFrame(analysis['summary_stats']))

                    # Add AI Analysis Option
                    if st.button("Generate AI Insights", type="primary"):
                        with st.spinner("Generating AI insights..."):
                            ai_insights = analyze_spreadsheet_with_openai(df)
                        st.subheader("AI-Generated Insights")
                        st.markdown(ai_insights)

                    # Export options
                    base_filename = uploaded_file.name.split('.')[0]
                    filename = f"analysis_{base_filename}"
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.download_button(
                            label="Download as Excel",
                            data=export_to_excel(df, filename),
                            file_name=f"{filename}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                    with col2:
                        st.download_button(
                            label="Download as CSV",
                            data=df.to_csv(index=False).encode('utf-8'),
                            file_name=f"{filename}.csv",
                            mime="text/csv",
                            use_container_width=True
                        )
                    with col3:
                        st.download_button(
                            label="Download Analysis as JSON",
                            data=export_to_json(json.dumps(analysis, indent=2), filename),
                            file_name=f"{filename}_analysis.json",
                            mime="application/json",
                            use_container_width=True
                        )
                else:
                    st.error(analysis)  # Error message from process_spreadsheet
            else:
                with st.spinner("Analyzing document..."):
                    extracted_text, extraction_log, extraction_method = summarize_extract_text_smart(uploaded_file, file_extension)

                with st.expander("Extraction Process Log", expanded=False):
                    for log_entry in extraction_log:
                        st.write(log_entry)

                if extracted_text:
                    method_info = {
                        "first_page": "First Page Content",
                        "index_pages": "Table of Contents/Index",
                        "first_page_fallback": "First Page (Limited Content)"
                    }
                    st.info(f"**Content Source:** {method_info.get(extraction_method, 'Unknown')}")

                    st.subheader("Extracted Text")
                    word_count = len(re.findall(r'\b\w+\b', extracted_text))
                    st.info(f"Extracted {len(extracted_text)} characters ({word_count} words)")

                    with st.expander("View extracted text", expanded=False):
                        st.text_area("Extracted Content", extracted_text, height=300, disabled=True)

                    if st.button("Generate AI Summary", type="primary", use_container_width=True):
                        if len(extracted_text.strip()) < 20:
                            st.warning("Text too short for summarization.")
                        else:
                            with st.spinner("Generating AI summary..."):
                                summary = summarize_text_with_openai(extracted_text, extraction_method)

                            st.subheader("AI-Generated Summary")
                            if summary and not summary.startswith("API Error"):
                                bullet_markers = ['•', '-', '*']
                                formatted_summary = summary
                                for marker in bullet_markers:
                                    formatted_summary = formatted_summary.replace(f'{marker} ', f'\n{marker} ')
                                lines = [line.strip() for line in formatted_summary.split('\n') if line.strip()]
                                formatted_summary = '\n'.join([f"{line}" for line in lines if line.startswith(tuple(bullet_markers))])
                                st.markdown(formatted_summary)
                            else:
                                st.error(summary)

                            if summary and not summary.startswith("API Error"):
                                method_suffix = "_index" if extraction_method == "index_pages" else "_firstpage"
                                base_filename = uploaded_file.name.split('.')[0]
                                filename = f"summary_{base_filename}{method_suffix}"
                                download_content = f"Document: {uploaded_file.name}\nContent Source: {method_info.get(extraction_method, 'Unknown')}\nSUMMARY:\n" + summary
                                
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.download_button(
                                        label="Download as Text",
                                        data=download_content,
                                        file_name=f"{filename}.txt",
                                        mime="text/plain",
                                        use_container_width=True
                                    )
                                with col2:
                                    st.download_button(
                                        label="Download as Word",
                                        data=export_to_word(download_content, filename),
                                        file_name=f"{filename}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        use_container_width=True
                                    )
                                with col3:
                                    st.download_button(
                                        label="Download as PDF",
                                        data=export_to_pdf(download_content, filename),
                                        file_name=f"{filename}.pdf",
                                        mime="application/pdf",
                                        use_container_width=True
                                    )
                                with col4:
                                    st.download_button(
                                        label="Download as JSON",
                                        data=export_to_json(download_content, filename),
                                        file_name=f"{filename}.json",
                                        mime="application/json",
                                        use_container_width=True
                                    )
                else:
                    st.error("Unable to extract readable text. Try a different file.")

    elif mode == "Advanced Analyzer":
        st.header("Advanced Text Analyzer (OCR Based)")
        st.markdown("Upload document or image and ask questions about content, tables, and data")

        # Initialize session state for analyzer
        if 'analyzer_vector_store' not in st.session_state:
            st.session_state.analyzer_vector_store = None
        if 'analyzer_qa_chain' not in st.session_state:
            st.session_state.analyzer_qa_chain = None
        if 'analyzer_processed_file' not in st.session_state:
            st.session_state.analyzer_processed_file = None
        if 'analyzer_document_type' not in st.session_state:
            st.session_state.analyzer_document_type = None

        _, tesseract_available = get_tesseract()
        use_ocr = st.checkbox("Enable OCR (PDF and Images)", value=tesseract_available, 
                              disabled=not tesseract_available, key="analyzer_ocr")

        uploaded_file = st.file_uploader("Choose PDF, Word, or Image document", 
                                       type=["pdf", "docx", "jpg", "png", "tiff"], 
                                       key="analyzer_uploader")

        if uploaded_file:
            file_extension = uploaded_file.name.split(".")[-1].lower()
            file_changed = (st.session_state.analyzer_processed_file != uploaded_file.name)

            if file_changed or not st.session_state.analyzer_vector_store:
                st.session_state.analyzer_vector_store = None
                st.session_state.analyzer_qa_chain = None
                st.session_state.analyzer_processed_file = uploaded_file.name

                # Determine and store document type
                st.session_state.analyzer_document_type = analyzer_determine_document_type(uploaded_file, file_extension, use_ocr)

                with st.spinner("Extracting text..."):
                    if file_extension == 'pdf':
                        text = analyzer_extract_text_from_pdf(uploaded_file)
                        ocr_text = ""
                        if use_ocr and tesseract_available:
                            with st.spinner("Running OCR..."):
                                ocr_text = analyzer_extract_images_and_ocr(uploaded_file)
                        elif use_ocr and not tesseract_available:
                            st.warning("OCR is not available due to missing Tesseract dependencies. Continuing with text extraction only.")
                    elif file_extension == 'docx':
                        text = analyzer_extract_text_from_docx(uploaded_file)
                        ocr_text = ""  # OCR not applicable for Word documents
                        if use_ocr:
                            st.warning("OCR is only supported for PDF and image files. Skipping OCR for Word document.")
                    elif file_extension in ['jpg', 'png', 'tiff']:
                        text = ""
                        ocr_text = analyzer_extract_text_from_image(uploaded_file)
                        if not use_ocr:
                            st.warning("OCR required for image files. Enable OCR to process.")
                    else:
                        st.error("Unsupported file type. Please upload a PDF, Word, or image document.")
                        text = None
                        ocr_text = ""

                full_text = text or ""
                if ocr_text:
                    full_text += "\n--- OCR CONTENT ---\n" + ocr_text

                if full_text:
                    st.success(f"Extracted {len(full_text):,} characters")

                    with st.spinner("Creating chunks..."):
                        text_splitter = analyzer_get_text_splitter()(
                            chunk_size=800,
                            chunk_overlap=100,
                            length_function=len
                        )
                        text_chunks = text_splitter.split_text(full_text)
                        st.info(f"Created {len(text_chunks)} chunks")

                    with st.spinner("Creating vector store..."):
                        st.session_state.analyzer_vector_store = analyzer_create_vector_store(text_chunks)

                    if st.session_state.analyzer_vector_store:
                        st.success("Ready for questions!")
                        st.session_state.analyzer_qa_chain = analyzer_create_qa_chain(
                            st.session_state.analyzer_vector_store, api_key
                        )

            # Display document type persistently
            if st.session_state.analyzer_document_type:
                st.markdown("### Document Type:")
                st.markdown(f"**{st.session_state.analyzer_document_type}**")

            if st.session_state.analyzer_vector_store and st.session_state.analyzer_qa_chain:
                st.markdown("---")
                st.subheader("Ask Questions")

                col1, col2, col3 = st.columns(3)
                with col1:
                    if st.button("Data", key="analyzer_data_btn"):
                        question = "What numerical data or statistics are mentioned?"
                        st.session_state.analyzer_current_question = question
                with col2:
                    if st.button("Key Points", key="analyzer_keypoints_btn"):
                        question = "What are the main findings or conclusions?"
                        st.session_state.analyzer_current_question = question

                question = st.text_input(
                    "Your question:",
                    value=st.session_state.get('analyzer_current_question', ''),
                    placeholder="Ask about content, tables, or specific information...",
                    key="analyzer_question_input"
                )

                if question:
                    with st.spinner("Processing..."):
                        try:
                            response = st.session_state.analyzer_qa_chain.invoke({"query": question})
                            answer = response["result"]
                            st.markdown("Answer:")
                            st.write(answer)
                            
                            if response.get("source_documents"):
                                with st.expander("Sources"):
                                    for i, doc in enumerate(response["source_documents"][:3]):
                                        st.markdown(f"**Source {i + 1}:**")
                                        content = doc.page_content
                                        st.text(content[:500] + "..." if len(content) > 500 else content)
                                        st.markdown("---")
                                
                                # Export options for answers
                                base_filename = uploaded_file.name.split('.')[0]
                                filename = f"answer_{base_filename}"
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.download_button(
                                        label="Download as Text",
                                        data=answer,
                                        file_name=f"{filename}.txt",
                                        mime="text/plain",
                                        use_container_width=True
                                    )
                                with col2:
                                    st.download_button(
                                        label="Download as Word",
                                        data=export_to_word(answer, filename),
                                        file_name=f"{filename}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        use_container_width=True
                                    )
                                with col3:
                                    st.download_button(
                                        label="Download as PDF",
                                        data=export_to_pdf(answer, filename),
                                        file_name=f"{filename}.pdf",
                                        mime="application/pdf",
                                        use_container_width=True
                                    )
                                with col4:
                                    st.download_button(
                                        label="Download as JSON",
                                        data=export_to_json(answer, filename),
                                        file_name=f"{filename}.json",
                                        mime="application/json",
                                        use_container_width=True
                                    )
                        except Exception as e:
                            st.error(f"Error: {str(e)}")

        else:
            st.info("Upload a PDF, Word, or image document to start")
            st.markdown("### Features:")
            st.markdown("""
            - **Text Extraction** - PyPDF2 for PDFs, python-docx for Word documents
            - **Image Support** - OCR for JPG, PNG, TIFF files
            - **OCR Support** - Tesseract for images, tables, and handwritten text (PDF and images)
            - **Smart Q&A** - OpenAI GPT-4o-mini with document analysis
            - **Data Analysis** - Handles reports, research papers, and images
            - **Export Options** - Word, PDF, JSON formats
            - **Fast Processing** - Optimized for performance
            """)

if __name__ == "__main__":
    main()
